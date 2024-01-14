import argparse
import csv
import os
import re
import sys
import time
import json
from pprint import pprint
import subprocess
import openai
from PIL import Image
import pytesseract
import PyPDF2
from docx import Document
import openpyxl
import pandas as pd


# globals

verbose = False

CREDS = 'credentials.json'
THREADS_CSV = 'threads.csv'
EXTRACTION_PERCENT = 10 # % of leading text to extract from files

asst_name = "File Renamer Assistant"
asst_instructions="""You help users rename files by generating a concise and descriptive new name based on the file text content 
    and calling rename_file function.
    """
asst_model="gpt-4-1106-preview" # models: gpt-3.5-turbo, gpt-4-1106-preview

orig_file_names = {}
dir_to_rename = None

# assistant functions

def rename_uploaded_file(old_name, new_name):
    new_dir = f'{dir_to_rename}renamed/'
    if dir_to_rename and not os.path.exists(new_dir):
        os.makedirs(new_dir)

    try: # replace uploaded f_id with old file name
        old_name = orig_file_names[old_name]
    except KeyError:
        pass
    if verbose:
        print(f'Renaming {old_name} to {new_name}')
    command = ['cp', f'{dir_to_rename}{old_name}', f'{new_dir}{new_name}']
    result = subprocess.run(command, capture_output=True, text=True)
    return result.stdout if result.returncode == 0 else result.stderr

def rename_file(old_name, new_name):
    new_dir = f'{dir_to_rename}renamed/'
    if dir_to_rename and not os.path.exists(new_dir):
        os.makedirs(new_dir)
    if verbose:
        print(f'Renaming {old_name} to {new_name}')
    command = ['cp', f'{dir_to_rename}{old_name}', f'{new_dir}{new_name}']
    result = subprocess.run(command, capture_output=True, text=True)
    return result.stdout if result.returncode == 0 else result.stderr

# assistant function interfaces

rename_uploaded_file_interface = {
    "name": "rename_uploaded_file",
    "description": """Renames the input file by passing old_name and new_name VALID JSON arguments, 
        e.g. \"arguments\": \"{\"old_name\":\"old_name.txt\", \"new_name\":\"new_name.txt\"}\"""",
    "parameters": {
        "type": "object",
        "properties": {
            "old_name": {
                "type": "string",
                "description": "The old name of the file, including it's extension; e.g. old_name.ext, NOT file_########"
            },
            "new_name": {
                "type": "string",
                "description": "The new name of the file, including it's extension; e.g. new_name.ext"
            }},
        "required": ["old_name", "new_name"]
    },}

rename_file_interface = {
    "name": "rename_file",
    "description": "Renames the input file by passing in <new_name>",
    "parameters": {
        "type": "object",
        "properties": {
            "old_name": {
                "type": "string",
                "description": "The old name of the file, including it's extension; e.g. old_name.ext"
            },
            "new_name": {
                "type": "string",
                "description": "The new name of the file, including the same extension as old_name; e.g. new_descriptive_name.ext"
            }},
        "required": ["old_name", "new_name"]
    },}

# assistant tools
asst_tools=[#{"type": "code_interpreter"},
            #{"type": "retrieval"},
            #{"type": "function", "function": rename_uploaded_file_interface},
            {"type": "function", "function": rename_file_interface},
        ]

# init

def get_creds():
    with open(CREDS, 'r') as file:
        return json.load(file)
def get_asst_id():
    creds = get_creds()
    try:
        return creds['asst_id']
    except KeyError as e:
        print()
        return None

creds = get_creds()
client = openai.OpenAI(api_key=creds['openai_api_key'])


# funcs

def show_json(obj):
    pprint(json.loads(obj.model_dump_json()))

def pprint_thread(thread):
    print("# Thread", thread)
    messages = client.beta.threads.messages.list(thread_id=thread.id, order="asc")
    return pprint_msgs(messages)

def pprint_msgs(messages):
    if verbose: print("Messages:")
    
    msg_str = ""
    for m in messages:
        line = f"{m.role}: {m.content[0].text.value}"
        if verbose:
            print(line)
            msg_str += line + '\n'

    if not verbose:
        print(line) # last line
        msg_str += line + '\n'
    
    return msg_str

def create_assistant():
    asst_id = get_asst_id()
    if asst_id:
        print(f'Assistant already created: {asst_id}')
        sys.exit(1)

    response = client.beta.assistants.create(name=asst_name, instructions=asst_instructions, model=asst_model, tools=asst_tools)
    # save assistant id to credentials
    asst_id = {'asst_id': f'{response.id}'}
    with open(CREDS, 'r+') as file:
        data = json.load(file)
        data.update(asst_id)
        file.seek(0)
        json.dump(data, file, indent=4)
        file.truncate()
    print(f'Created Assistant: {response.id}')
    print(f'asst_id saved to credentials.json')
    return response

def update_assistant():
    asst_id = get_asst_id()
    client.beta.assistants.update(asst_id, name=asst_name, instructions=asst_instructions, model=asst_model, tools=asst_tools)
    
    print(f"Updated Assistant: {asst_id}")

def file_delete(f_id):
    response = client.files.delete(f_id)
    print(response)

def upload_file_for_asst(file_path):
    # check if file exists
    if not os.path.exists(file_path):
        print(f'Error: file {file_path} does not exist')
        sys.exit(1)

    asst_id = get_asst_id()

    response = client.files.create(
        file=open(file_path, "rb"),
        purpose="assistants")
    try:
        asst_file = client.beta.assistants.files.create(
        assistant_id=asst_id,
        file_id=response.id
        )
        print(asst_file)
    except Exception as e:
        print(e)
        file_delete(response.id)
        return None
    
    return response.id

def get_slice_size(total):
    slice_len = int(total * EXTRACTION_PERCENT / 100) # Calculate the slice length as integer

    if not slice_len: # slice_len = 0
        raise ValueError('Error: No text extracted. EXTRACTION_PERCENT needs to be increased to rename this file.')
    return slice_len

def extract_text_from_file(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension in ['.txt']:
        with open(file_path, 'r') as file:
            content = file.readlines()
            return ''.join(content[:get_slice_size(len(content))])

    elif file_extension in ['.pdf']:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            total_pages = len(reader.pages)
            text = [reader.pages[i].extract_text() for i in range(get_slice_size(total_pages))]
            return "\n".join(text)

    elif file_extension in ['.docx']:
        doc = Document(file_path)
        total_paragraphs = len(doc.paragraphs)
        return "\n".join(para.text for para in doc.paragraphs[:get_slice_size(total_paragraphs)])

    elif file_extension in ['.xlsx']:
        wb = openpyxl.load_workbook(file_path)
        sheet = wb.active
        total_rows = len(list(sheet.rows))
        return "\n".join(str(cell.value) for row in list(sheet.rows)[:get_slice_size(total_rows)] for cell in row)

    elif file_extension in ['.csv']:
        # Reading the CSV file using pandas and converting a portion of it to string
        df = pd.read_csv(file_path)
        slice_length = get_slice_size(len(df))
        return df.head(slice_length).to_string(index=False)

    elif file_extension in ['.jpg', '.jpeg', '.png']:
        img = Image.open(file_path)
        return pytesseract.image_to_string(img)
    else:
        raise ValueError("Error: Unsupported file format")

def rename_files(dir_path):
    global dir_to_rename
    dir_to_rename = dir_path

    file_list = os.listdir(dir_path)
    for file in file_list:
        f_path = dir_path + file
        
        # skip subdirectories
        if not os.path.isfile(f_path):
            continue

        try:
            if verbose:
                print(f'Extracting text from {file}.')
            file_text = None
            file_text = extract_text_from_file(f_path)
            if verbose:
                print(f'Extracted text: {file_text}')
        except Exception as e:
            print(f'Skipping {file}. {e}')
            continue

        query_last_thread(f"""Generate a concise, meaningful file name for a file ({file}) containing the following text content: {file_text}\n
            Then rename it using rename_file_interface.""")

def create_thread():
    thread = client.beta.threads.create()
    # create threads.csv if not exists
    if not os.path.exists(THREADS_CSV):
        # write header to csv
        with open(THREADS_CSV, 'w') as f:
            f.write('thread_id\n')
    # write thread id to csv
    with open(THREADS_CSV, 'a') as f:
        f.write(thread.id + '\n')
    
    return thread

def create_run(thread, user_input):
    run = submit_message(get_asst_id(), thread, user_input)
    return run

def submit_message(assistant_id, thread, user_message):
    client.beta.threads.messages.create(
        thread_id=thread.id, role="user", content=user_message
    )
    return client.beta.threads.runs.create(
        thread_id=thread.id,
        assistant_id=assistant_id,
    )

def wait_on_run(run, thread):
    while run.status == "queued" or run.status == "in_progress":
        run = client.beta.threads.runs.retrieve(
            thread_id=thread.id,
            run_id=run.id,
        )
        time.sleep(1)
    return run

def get_response(thread):
    return client.beta.threads.messages.list(thread_id=thread.id, order="asc")

def thread_get(thread_id):
    if thread_id == "new":
        thread = create_thread()
    else:
        thread = client.beta.threads.retrieve(thread_id)
    return thread

def get_last_thread_id():
    with open(THREADS_CSV, 'r', newline='') as f:
        reader = csv.reader(f)
        last_line = None
        for line in reader:
            last_line = line
        if last_line and last_line[0] == 'thread_id':
            last_line = None
        return last_line[0] if last_line else None

def steps_get(thread, run):
    # get steps
    run_steps = client.beta.threads.runs.steps.list(thread_id=thread.id, run_id=run.id, order="asc")
    if verbose:
        for step in run_steps.data:
            step_details = step.step_details
            print(json.dumps(show_json(step_details), indent=4))
    return run_steps

def thread_delete(thread_id):
    client.beta.threads.delete(thread_id)
    with open(THREADS_CSV, 'r', newline='') as file:
        reader = csv.reader(file)
        lines = [line for line in reader if line[0] != thread_id]

    with open(THREADS_CSV, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerows(lines)
        
    if verbose:
        print("Deleted thread:", thread_id)

def call_tool(run, thread):
    # Extract single tool call
    tool_calls = run.required_action.submit_tool_outputs.tool_calls

    tool_outputs = []
    for tool_call in tool_calls:
        if verbose:
            print(f'Tool call arguments: {tool_call.function.arguments}')
        name = tool_call.function.name
        arguments = json.loads(tool_call.function.arguments)

        if verbose:
            print("Function Name:", name)
            print("Function Arguments:", arguments)

    if name == "rename_uploaded_file":
        responses = rename_uploaded_file(arguments["old_name"], arguments["new_name"])
    elif name == "rename_file":
        responses = rename_file(arguments["old_name"], arguments["new_name"])
    
    tool_outputs.append({"tool_call_id": tool_call.id, "output": json.dumps(responses)})
    # submit tool outputs
    run = client.beta.threads.runs.submit_tool_outputs(thread_id=thread.id, run_id=run.id, tool_outputs=tool_outputs)

    return wait_on_run(run, thread)

def query(user_input, thread=None):

    if not thread:
        # create thread and run
        thread = create_thread()
        if verbose:
            print("Thread ID:", thread.id)

    # create new run
    try:
        run = create_run(thread, user_input)
    except openai.BadRequestError as e:
        error_message = str(e)

        # Handle error: 'Cannot add messages to <thread_> while a run <run_> is active'

        # Extracting thread and run IDs using regex
        thread_id = None
        run_id = None
        pattern = r"(thread_[\w\d]+)|(run_[\w\d]+)"
        matches = re.findall(pattern, error_message)
        for match in matches:
            if match[0]:  # corresponds to thread ID
                thread_id = match[0]
            if match[1]:  # corresponds to run ID
                run_id = match[1]

        # Removing thread and run strings from the error message
        cleaned_error_message = re.sub(pattern, "", error_message).strip()
        # Check if error is due to an active run
        if "Can't add messages to  while a run  is active" in cleaned_error_message:
            if verbose:
                print(f"Previous run {run_id} still active. Cancelling it...")
            run = client.beta.threads.runs.cancel(thread_id=thread_id, run_id=run_id)
            # wait for run cancellation
            run = wait_on_run(run, thread)
            
            # create new run
            if verbose:
                print("Creating new run...")
            run = create_run(thread, user_input)
        else:
            raise e # other BadRequestErrors

    if verbose:
        print("Run ID:", run.id)

    # wait for run completion
    run = wait_on_run(run, thread)
    
    if verbose:
        print("Run status: ", run.status)
    if run.status == 'failed':
        print('Run status: failed.')
        print(f'Last Error: {run.last_error}')

    while run.status == "requires_action":
        run = call_tool(run, thread)
    

    if run.status == "completed":
        response = get_response(thread)
        return pprint_msgs(response)

def query_last_thread(q):
    lt_id = get_last_thread_id()
    #print(f'last_thread={lt_id}')
    if not lt_id:
        if verbose:
            print(f'No threads in {THREADS_CSV}')
        thread_get('new')
        lt_id = get_last_thread_id()
        if verbose:
            print(f'Created new thread: {lt_id}')
    return query(q, thread_get(lt_id))

def main(args):
    global verbose
    verbose = args.verbose

    if args.asst_create:
        asst = create_assistant()
    elif args.asst_update:
        update_assistant()
    elif args.asst_file_upload:
        upload_file_for_asst(args.asst_file_upload)
    elif args.files_list:
        show_json(client.files.list())
    elif args.file_delete:
        file_delete(args.file_delete)
    elif args.files_rename:
        rename_files(args.files_rename)
    elif args.query_new:
        query(args.query_new, None)
    elif args.query_last_thread:
        query_last_thread(args.query_last_thread)
    elif args.thread_get:
        thread = thread_get(args.thread_get)
        return thread
    elif args.thread_delete:
        thread_delete(args.thread_delete)
    elif args.steps_get:
        thread = client.beta.threads.retrieve(args.steps_get[0])
        run = client.beta.threads.runs.retrieve(thread_id=thread.id, run_id=args.steps_get[1])
        return steps_get(thread, run)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='OpenAI Assistant to rename directory files in a given directory.')
    parser.add_argument('--asst_create', '-ac', action='store_true', help='Create the Assistant')
    parser.add_argument('--asst_update', '-au', action='store_true', help='Update the Assistant')
    parser.add_argument('--asst_file_upload', '-afu', type=str, help='Upload file for Assistant to retrieve; input: file_path')
    parser.add_argument('--files_list', '-fl', action='store_true', help='List organization\'s files')
    parser.add_argument('--file_delete', '-fd', type=str, help='Delete file; input: file_id')
    parser.add_argument('--files_rename', '-fr', type=str, help='Rename all files in directory; input: dir_path')
    parser.add_argument('--query_new', '-qn', type=str, help='Create new thread and run query')
    parser.add_argument('--query_last_thread', '-qlt', type=str, help='Append query to last thread')
    parser.add_argument('--thread_get', '-tg', type=str, help='Get a thread; input: thread_id, "new"')
    parser.add_argument('--thread_delete', '-td', type=str, help='Delete a thread; input: thread_id')
    parser.add_argument('--steps_get', '-sg', nargs=2, help='Get run steps; input: thread_id, run_id')
    parser.add_argument('--verbose', '-v', action='store_true', help='Enable verbose output')

    main(parser.parse_args())